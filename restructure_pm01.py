"""
Restructure УМКД ПМ01 from flat format to LO-based format matching ПМ10 example.
Output: УМКД_ПМ01_Structured.docx
"""

import os
import copy
from docx import Document
from docx.shared import Pt, Cm, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from lxml import etree

# ── paths ────────────────────────────────────────────────────────────────────
FOLDER = r'C:\Users\Dias\Desktop\УМКД'
_docx = sorted([f for f in os.listdir(FOLDER) if f.endswith('.docx') and 'Structured' not in f])
# _docx[0] = Пример ПМ10,  _docx[1] = УМКД_ПМ01
PM01_PATH   = os.path.join(FOLDER, _docx[1])
OUTPUT_PATH = os.path.join(FOLDER, 'УМКД_ПМ01_Structured.docx')

FONT      = 'Times New Roman'
SZ_TITLE  = 14   # pt  – cover page
SZ_BODY   = 12   # pt  – all body text  (matches PM10)
SZ_CODE   = 11   # pt  – code blocks (Courier New)

# ── LO definitions ────────────────────────────────────────────────────────────
LO_LIST = [
    {
        'id': '1.1',
        'title': 'Comply with occupational and electrical safety requirements in information-service enterprises',
        'weight': '8%',
        'lecture_nums': [1, 2, 3, 4],
        'practical_label': 'Practical class',
        'practical_nums': [1, 2],
        'cal_rows': [1, 2, 3, 4, 27, 28],
        'glossary_terms': ['ESD', 'ESD wrist strap', 'Grounding', 'Static electricity'],
        'hours': '22 hours',
    },
    {
        'id': '1.2',
        'title': 'Determine the operability of computer systems at the circuit level',
        'weight': '10%',
        'lecture_nums': [5, 6, 7, 8, 9, 10, 11],
        'practical_label': 'Practical class',
        'practical_nums': [3, 4, 5, 6],
        'cal_rows': [5, 6, 7, 8, 9, 10, 11, 29, 30, 31, 32],
        'glossary_terms': ['Bus', 'Cache', 'Chipset', 'Connector', 'CPU',
                           'Firmware', 'Form factor', 'GPU', 'Heat sink', 'Jumper',
                           'Motherboard', 'Multimeter', 'Northbridge', 'Socket',
                           'Southbridge', 'Voltage', 'Wattage'],
        'hours': '44 hours',
    },
    {
        'id': '1.3',
        'title': 'Operate and monitor computers and peripheral devices',
        'weight': '8%',
        'lecture_nums': [12, 13, 14, 15],
        'practical_label': 'Laboratory work',
        'practical_nums': [1, 2],
        'cal_rows': [12, 13, 14, 15, 33, 34],
        'glossary_terms': ['Driver', 'HDD', 'Overheating', 'Peripheral',
                           'RAM', 'SATA', 'SMART', 'SSD', 'Thermal throttling', 'USB'],
        'hours': '22 hours',
    },
    {
        'id': '1.4',
        'title': 'Diagnose and eliminate hardware faults',
        'weight': '10%',
        'lecture_nums': [16, 17, 18, 19],
        'practical_label': 'Laboratory work',
        'practical_nums': [3, 4],
        'cal_rows': [16, 17, 18, 19, 35, 36],
        'glossary_terms': ['Beep code', 'Diagnostics', 'POST', 'Troubleshooting'],
        'hours': '22 hours',
    },
    {
        'id': '1.5',
        'title': 'Restore equipment after failures and prevent them',
        'weight': '8%',
        'lecture_nums': [20, 21, 22, 23],
        'practical_label': 'Laboratory work',
        'practical_nums': [5, 6],
        'cal_rows': [20, 21, 22, 23, 37, 38],
        'glossary_terms': ['Backup', 'BIOS', 'RAID', 'SMPS',
                           'Thermal paste', 'UEFI', 'UPS'],
        'hours': '22 hours',
    },
    {
        'id': '1.6',
        'title': 'Replace consumables used in computing and office equipment',
        'weight': '8%',
        'lecture_nums': [24, 25, 26],
        'practical_label': 'Industrial training',
        'practical_nums': list(range(1, 7)),
        'cal_rows': [24, 25, 26, 39, 40, 41, 42, 43, 44],
        'glossary_terms': ['Cartridge', 'Toner'],
        'hours': '66 hours',
    },
    {
        'id': '1.7',
        'title': 'Manage the internal and external devices of a personal computer',
        'weight': '8%',
        'lecture_nums': [],
        'practical_label': 'Industrial training',
        'practical_nums': list(range(7, 19)),
        'cal_rows': list(range(45, 57)),
        'glossary_terms': ['PSU', 'Connector', 'USB', 'Form factor', 'Socket'],
        'hours': '54 hours',
    },
]

# Software list for every LO (same across all)
SOFTWARE_LIST = [
    '– Multimedia projector and interactive (smart) board.',
    '– PowerPoint presentations for the module topics.',
    '– Disassembled training PC stands; sets of motherboards, power supplies, RAM modules, storage drives and graphics adapters.',
    '– Peripheral and office equipment: printers, scanners, multifunction devices (MFPs).',
    '– Measuring instruments: multimeters, POST-cards, power-supply and cable testers.',
    '– Anti-static wrist straps and mats; PC assembly and repair tool kits; soldering stations.',
    '– Hardware diagnostic software: AIDA64, MemTest86, Victoria, CrystalDiskInfo, HWMonitor.',
    '– Consumables of office equipment (cartridges, toner, ink, ribbons) for production training.',
    '– Learning Management System (LMS) for hosting materials and assignments.',
    '– Video lectures and educational videos on PC maintenance and repair.',
]

LITERATURE_CORE = [
    '1. Mueller S. Upgrading and Repairing PCs. — 21st ed. — Moscow: Williams, 2019. — 1376 p.',
    '2. Loginov M.D., Loginova T.A. Technical Maintenance of Computing Equipment: a study guide. — Moscow: BINOM, 2019. — 319 p.',
    '3. Partyka T.L., Popov I.I. Peripheral Devices of Computing Equipment: a study guide. — 3rd ed. — Moscow: FORUM; INFRA-M, 2018. — 432 p.',
    '4. Kolesnichenko O.V., Shishigin I.V. PC Hardware. — Saint Petersburg: BHV-Petersburg, 2017. — 800 p.',
]
LITERATURE_EXTRA = [
    '1. Tanenbaum A., Austin T. Structured Computer Organization. — 6th ed. — Saint Petersburg: Piter, 2018. — 816 p.',
    '2. Guk M. IBM PC Hardware. Encyclopedia. — 3rd ed. — Saint Petersburg: Piter, 2016. — 1072 p.',
]
LITERATURE_WEB = [
    '1. intel.com — official technical materials and specifications of Intel processors and chipsets.',
    '2. hardwareinside.ru — articles and guides on PC hardware diagnostics and repair.',
    '3. ixbt.com — reviews, tests and technical descriptions of computer components and peripherals.',
    '4. habr.com/ru/hub/hardware/ — publications on circuitry, repair and maintenance of computing equipment.',
]

# ═══════════════════════════════════════════════════════════════════════════
#  EXTRACT PM01 content
# ═══════════════════════════════════════════════════════════════════════════

def extract_pm01():
    doc = Document(PM01_PATH)

    # map body elements
    body = doc.element.body
    all_elems = []
    para_map = {p._element: p for p in doc.paragraphs}
    tbl_map  = {t._element: t for t in doc.tables}
    for child in body:
        tag = child.tag.split('}')[-1]
        if tag == 'p' and child in para_map:
            all_elems.append(('para', para_map[child]))
        elif tag == 'tbl' and child in tbl_map:
            all_elems.append(('table', tbl_map[child]))

    tables = doc.tables

    # ── Calendar rows (table index 2) ──
    cal_table = tables[2]
    cal_rows = {}  # row_number (int) -> [no, topic, date, activity]
    for ri, row in enumerate(cal_table.rows):
        cells = [c.text.strip() for c in row.cells]
        if ri == 0:
            continue  # skip header
        try:
            num = int(cells[0].rstrip('.'))
            cal_rows[num] = cells
        except ValueError:
            pass

    # ── Glossary rows (table index 3) ──
    gloss_table = tables[3]
    gloss = {}  # term -> definition
    for ri, row in enumerate(gloss_table.rows):
        if ri == 0:
            continue
        cells = [c.text.strip() for c in row.cells]
        if cells[0]:
            gloss[cells[0]] = cells[1] if len(cells) > 1 else ''

    # ── Full text sections: find section start indices in all_elems ──
    def find_elem(text_fragment, from_idx=0):
        for i, (et, el) in enumerate(all_elems[from_idx:], start=from_idx):
            if et == 'para' and text_fragment in el.text:
                return i
        return -1

    intro_start    = find_elem('INTRODUCTION')
    lecture_start  = find_elem('III. LECTURE COMPLEX')
    practical_start= find_elem('IV. PRACTICAL CLASSES')
    lab_start      = find_elem('V. LABORATORY CLASSES')
    industrial_start= find_elem('VI. INDUSTRIAL TRAINING')
    indwork_start  = find_elem('VII. INDEPENDENT WORK')
    software_start = find_elem('VIII. LIST OF SOFTWARE')
    ref_start      = find_elem('IX. LIST OF REFERENCES')

    # ── Introduction paragraphs: only the body text, NOT headers/table ──
    # In PM01: INTRODUCTION header → 6 body paras → "Module assessment weights" header → table
    weights_idx = find_elem('Module assessment weights', intro_start)
    if weights_idx < 0:
        weights_idx = lecture_start
    intro_paras = []
    # start from intro_start+1 to skip the INTRODUCTION header itself
    for et, el in all_elems[intro_start + 1: weights_idx]:
        if et == 'para':
            intro_paras.append(el)

    # ── Lectures: group by lecture number ──
    lectures = {}   # lecture_num -> list of (etype, elem)
    cur_lec = None
    for et, el in all_elems[lecture_start:practical_start]:
        if et == 'para':
            txt = el.text.strip()
            if txt.startswith('Lecture No.'):
                try:
                    cur_lec = int(txt.split('No.')[-1].strip())
                    lectures[cur_lec] = []
                except ValueError:
                    pass
            if cur_lec is not None:
                lectures[cur_lec].append((et, el))
        elif et == 'table' and cur_lec is not None:
            lectures[cur_lec].append((et, el))

    # ── Practical classes: group by number ──
    practical_classes = {}
    cur_pc = None
    for et, el in all_elems[practical_start:lab_start]:
        if et == 'para':
            txt = el.text.strip()
            if txt.startswith('Practical Class No.'):
                try:
                    cur_pc = int(txt.split('No.')[-1].strip())
                    practical_classes[cur_pc] = []
                except ValueError:
                    pass
            if cur_pc is not None:
                practical_classes[cur_pc].append((et, el))
        elif et == 'table' and cur_pc is not None:
            practical_classes[cur_pc].append((et, el))

    # ── Laboratory works: group by number ──
    lab_works = {}
    cur_lb = None
    for et, el in all_elems[lab_start:industrial_start]:
        if et == 'para':
            txt = el.text.strip()
            if txt.startswith('Laboratory Work No.'):
                try:
                    cur_lb = int(txt.split('No.')[-1].strip())
                    lab_works[cur_lb] = []
                except ValueError:
                    pass
            if cur_lb is not None:
                lab_works[cur_lb].append((et, el))
        elif et == 'table' and cur_lb is not None:
            lab_works[cur_lb].append((et, el))

    # ── Industrial training tasks: group by task number ──
    ind_tasks = {}   # task_num -> list of (etype, elem)
    cur_task = None
    for et, el in all_elems[industrial_start:indwork_start]:
        if et == 'para':
            txt = el.text.strip()
            if txt.startswith('Industrial Training Task No.'):
                try:
                    cur_task = int(txt.split('No.')[-1].strip())
                    ind_tasks[cur_task] = []
                except ValueError:
                    pass
            if cur_task is not None:
                ind_tasks[cur_task].append((et, el))

    # ── Independent work (Self-Study) ──
    self_study = []
    for et, el in all_elems[indwork_start:software_start]:
        self_study.append((et, el))

    return {
        'intro_paras': intro_paras,
        'cal_rows': cal_rows,
        'gloss': gloss,
        'lectures': lectures,
        'practical_classes': practical_classes,
        'lab_works': lab_works,
        'ind_tasks': ind_tasks,
        'self_study': self_study,
    }


# ═══════════════════════════════════════════════════════════════════════════
#  FORMATTING HELPERS
# ═══════════════════════════════════════════════════════════════════════════

def _set_run(run, text, bold=False, size=SZ_BODY, font=FONT, italic=False):
    run.text = text
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic

def _set_para_fmt(para, align=WD_ALIGN_PARAGRAPH.JUSTIFY,
                  space_before=0, space_after=0, line_spacing=None):
    pf = para.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after  = Pt(space_after)
    if line_spacing is not None:
        pf.line_spacing = Pt(line_spacing)

def add_para(doc, text='', bold=False, size=SZ_BODY, font=FONT,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY,
             space_before=0, space_after=0):
    p = doc.add_paragraph()
    _set_para_fmt(p, align=align, space_before=space_before, space_after=space_after)
    if text:
        r = p.add_run(text)
        _set_run(r, text, bold=bold, size=size, font=font)
    return p

def add_page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(0)
    run = p.add_run()
    run.add_break(docx_break_type())

def docx_break_type():
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    return _BreakWrapper(br)

class _BreakWrapper:
    """Helper so we can pass break XML to run.add_break."""
    def __init__(self, br_elem):
        self._br = br_elem

def _real_page_break(doc):
    """Insert a proper page break paragraph."""
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)

def set_margins(doc):
    """Set PM10-compatible margins on all sections."""
    for section in doc.sections:
        sectPr = section._sectPr
        pgMar = sectPr.find(qn('w:pgMar'))
        if pgMar is None:
            pgMar = OxmlElement('w:pgMar')
            sectPr.append(pgMar)
        # top=2cm, bottom=2cm, left=2.5cm, right=1.5cm  (in twips)
        pgMar.set(qn('w:top'),    '1134')
        pgMar.set(qn('w:bottom'), '1134')
        pgMar.set(qn('w:left'),   '1418')
        pgMar.set(qn('w:right'),  '851')

def _cell_para(cell, text, bold=False, size=SZ_BODY,
               align=WD_ALIGN_PARAGRAPH.LEFT):
    """Clear cell and add single paragraph."""
    cell.paragraphs[0].clear()
    p = cell.paragraphs[0]
    _set_para_fmt(p, align=align)
    r = p.add_run(text)
    _set_run(r, text, bold=bold, size=size)
    return p

def _set_table_borders(table):
    """Apply thin borders to every cell."""
    tbl = table._tbl
    tblPr = tbl.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        tbl.insert(0, tblPr)
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'),  '4')
        el.set(qn('w:space'),'0')
        el.set(qn('w:color'),'000000')
        tblBorders.append(el)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)

def _set_col_widths(table, widths_cm):
    """Set column widths in cm."""
    for ri, row in enumerate(table.rows):
        for ci, cell in enumerate(row.cells):
            if ci < len(widths_cm):
                tc = cell._tc
                tcPr = tc.find(qn('w:tcPr'))
                if tcPr is None:
                    tcPr = OxmlElement('w:tcPr')
                    tc.insert(0, tcPr)
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                twips_val = int(widths_cm[ci] * 567)   # 1cm ≈ 567 twips
                tcW.set(qn('w:w'), str(twips_val))
                tcW.set(qn('w:type'), 'dxa')


def copy_para_to_doc(src_para, doc, force_size=SZ_BODY):
    """Copy a paragraph from PM01 into new doc, resizing TNR runs to force_size."""
    p = doc.add_paragraph()
    pf = src_para.paragraph_format
    npf = p.paragraph_format
    npf.alignment = pf.alignment
    try:
        npf.space_before = pf.space_before or Pt(0)
    except Exception:
        npf.space_before = Pt(0)
    try:
        npf.space_after = pf.space_after or Pt(0)
    except Exception:
        npf.space_after = Pt(0)

    for run in src_para.runs:
        nr = p.add_run(run.text)
        nr.bold   = run.bold
        nr.italic = run.italic
        nr.underline = run.underline
        fn = run.font.name or FONT
        nr.font.name = fn
        # resize: keep Courier New as-is, resize TNR 14→12
        if fn and fn.lower() == 'courier new':
            nr.font.size = Pt(SZ_CODE)
        else:
            nr.font.size = Pt(force_size)
    return p


def copy_table_to_doc(src_table, doc):
    """Deep-copy a table element from PM01 into doc."""
    new_tbl = copy.deepcopy(src_table._tbl)
    # Resize font in all runs
    for r_elem in new_tbl.findall('.//' + qn('w:r')):
        rPr = r_elem.find(qn('w:rPr'))
        if rPr is not None:
            sz_elem = rPr.find(qn('w:sz'))
            szCs_elem = rPr.find(qn('w:szCs'))
            fn_elem = rPr.find(qn('w:rFonts'))
            is_code = fn_elem is not None and 'Courier' in (fn_elem.get(qn('w:ascii')) or '')
            target = SZ_CODE if is_code else SZ_BODY
            half = str(target * 2)
            if sz_elem is None:
                sz_elem = OxmlElement('w:sz')
                rPr.append(sz_elem)
            sz_elem.set(qn('w:val'), half)
            if szCs_elem is None:
                szCs_elem = OxmlElement('w:szCs')
                rPr.append(szCs_elem)
            szCs_elem.set(qn('w:val'), half)
    doc.element.body.append(new_tbl)


# ═══════════════════════════════════════════════════════════════════════════
#  BUILD SECTIONS
# ═══════════════════════════════════════════════════════════════════════════

def build_title_page(doc):
    # College name
    add_para(doc, '«APEC PETROTECHNIC HIGHER COLLEGE» LLP',
             bold=True, size=SZ_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)

    # APPROVED table (2 cols: left blank, right approved block)
    tbl = doc.add_table(rows=1, cols=2)
    _set_table_borders(tbl)
    tbl.style = 'Table Grid'
    left_cell  = tbl.rows[0].cells[0]
    right_cell = tbl.rows[0].cells[1]
    _set_col_widths(tbl, [9.0, 8.0])
    left_cell.paragraphs[0].text = ''
    def _rc(text, bold=False, align=WD_ALIGN_PARAGRAPH.RIGHT):
        p = right_cell.add_paragraph()
        _set_para_fmt(p, align=align, space_after=0)
        r = p.add_run(text)
        _set_run(r, text, bold=bold, size=SZ_BODY)
    right_cell.paragraphs[0].clear()
    _rc('«APPROVED»', bold=True, align=WD_ALIGN_PARAGRAPH.RIGHT)
    _rc('Deputy director', align=WD_ALIGN_PARAGRAPH.RIGHT)
    _rc('for academic work', align=WD_ALIGN_PARAGRAPH.RIGHT)
    _rc('___________ Bissembayeva G.Y.', align=WD_ALIGN_PARAGRAPH.RIGHT)
    _rc('"____" ____________2024 year', align=WD_ALIGN_PARAGRAPH.RIGHT)

    for _ in range(3):
        add_para(doc)

    add_para(doc, 'Educational-methodological complex of the subject',
             bold=True, size=SZ_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    add_para(doc)
    add_para(doc, 'Module: PM 01 «Maintenance and configuration of computer hardware and software»',
             bold=True, size=SZ_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    add_para(doc, 'Specialty: 06130100 – «Software Engineering»',
             size=SZ_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER)

    for _ in range(6):
        add_para(doc)

    add_para(doc, 'Atyrau, 2025',
             bold=True, size=SZ_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER)


def build_approval_sheet(doc):
    _real_page_break(doc)
    add_para(doc, '«APEC PETROTECHNIC HIGHER COLLEGE» LLP',
             bold=True, size=SZ_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    add_para(doc, 'APPROVAL SHEET',
             bold=True, size=SZ_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    add_para(doc, 'Reviewed and approved at a meeting of the Methodological Council',
             size=SZ_BODY)
    add_para(doc, 'Protocol № ____ of «____» ____________ 2026 y.',
             size=SZ_BODY)
    add_para(doc)
    add_para(doc, 'Considered at the meeting of the CMC Subject-Cycle Commission «Software Engineering»',
             size=SZ_BODY)
    add_para(doc, 'Protocol № ____ of «____» ____________ 2026 y.',
             size=SZ_BODY)
    add_para(doc)
    add_para(doc, 'Developer: __________ D.M. Yermekov', size=SZ_BODY)


def build_contents(doc):
    _real_page_break(doc)
    add_para(doc, 'CONTENT', bold=True, size=SZ_BODY,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)

    def _toc_line(text, page, bold=False, indent=False):
        p = doc.add_paragraph()
        _set_para_fmt(p, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        if indent:
            p.paragraph_format.left_indent = Pt(24)
        r = p.add_run(text)
        _set_run(r, text, bold=bold, size=SZ_BODY)
        r2 = p.add_run(f'\t{page}')
        _set_run(r2, f'\t{page}', bold=bold, size=SZ_BODY)
        pPr = p._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        tab = OxmlElement('w:tab')
        tab.set(qn('w:val'),    'right')
        tab.set(qn('w:leader'), 'dot')
        tab.set(qn('w:pos'),    '9350')
        tabs.append(tab)
        pPr.append(tabs)

    _toc_line('INTRODUCTION', 4, bold=True)
    _toc_line('I. ACADEMIC CALENDAR OF MODULE', 5, bold=True)

    lo_pages = {'1.1': 6, '1.2': 14, '1.3': 30, '1.4': 40,
                '1.5': 52, '1.6': 64, '1.7': 78}
    roman = {0: 'II', 1: 'III', 2: 'IV', 3: 'V', 4: 'VI', 5: 'VII', 6: 'VIII'}

    for idx, lo in enumerate(LO_LIST):
        lid = lo['id']
        base = lo_pages[lid]
        _toc_line(f'{roman[idx]}. LO {lid} {lo["title"]}', base, bold=True)
        sub = 1
        _toc_line(f'{sub}. LO {lid} Academic calendar', base + 1, indent=True); sub += 1
        _toc_line(f'{sub}. LO {lid} Glossary', base + 3, indent=True); sub += 1
        if lo['lecture_nums']:
            _toc_line(f'{sub}. LO {lid} Lecture materials', base + 5, indent=True); sub += 1
        prac_label = lo['practical_label']
        if prac_label == 'Industrial training':
            _toc_line(f'{sub}. LO {lid} Industrial training materials', base + 15, indent=True)
        else:
            _toc_line(f'{sub}. LO {lid} Practical class materials', base + 15, indent=True)
        sub += 1
        _toc_line(f'{sub}. LO {lid} List of software and multimedia support for training sessions',
                  base + 20, indent=True); sub += 1
        _toc_line(f'{sub}. LO {lid} List of literature', base + 22, indent=True)

    _toc_line('IX. Self-Study Topics for the Module', 101, bold=True)
    _toc_line('X. Exam Questions', 102, bold=True)


def build_introduction(doc, data):
    _real_page_break(doc)
    add_para(doc, 'INTRODUCTION', bold=True, size=SZ_BODY,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    for para in data['intro_paras']:
        txt = para.text.strip()
        if not txt:
            add_para(doc)
            continue
        copy_para_to_doc(para, doc, force_size=SZ_BODY)


def build_weight_table(doc):
    _real_page_break(doc)
    add_para(doc, 'ACADEMIC CALENDAR OF MODULE', bold=True, size=SZ_BODY,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)

    # 3 columns: LO code | description | weight %
    nrows = 1 + len(LO_LIST) + 2  # header + LOs + empty row + intermediate
    tbl = doc.add_table(rows=nrows, cols=3)
    _set_table_borders(tbl)
    _set_col_widths(tbl, [2.5, 11.5, 2.5])

    # Header row
    _cell_para(tbl.rows[0].cells[0], 'PM 01', bold=True,
               align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    _cell_para(tbl.rows[0].cells[1],
               'Maintenance and configuration of computer hardware and software',
               bold=True, align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    _cell_para(tbl.rows[0].cells[2], 'Weight. %', bold=True,
               align=WD_ALIGN_PARAGRAPH.CENTER)

    for i, lo in enumerate(LO_LIST):
        row = tbl.rows[i + 1]
        _cell_para(row.cells[0], f'LO {lo["id"]}',
                   align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(row.cells[1], lo['title'],
                   align=WD_ALIGN_PARAGRAPH.JUSTIFY)
        _cell_para(row.cells[2], lo['weight'],
                   align=WD_ALIGN_PARAGRAPH.CENTER)

    # Empty row
    empty_row = tbl.rows[1 + len(LO_LIST)]
    for cell in empty_row.cells:
        cell.paragraphs[0].text = ''

    # Intermediate certification
    last_row = tbl.rows[-1]
    _cell_para(last_row.cells[0], '', align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_para(last_row.cells[1], 'Intermediate certification',
               align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    _cell_para(last_row.cells[2], '40%',
               align=WD_ALIGN_PARAGRAPH.CENTER)


def build_lo_calendar(doc, lo, data):
    add_para(doc, 'ACADEMIC CALENDAR',
             bold=True, size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    add_para(doc, f'LO {lo["id"]} {lo["title"]}', size=SZ_BODY)
    add_para(doc, 'Instructor: D.M. Yermekov', size=SZ_BODY)
    add_para(doc, 'Specialty: 06130100 – «Software Engineering»', size=SZ_BODY)
    add_para(doc, f'Total hours: {lo["hours"]}', size=SZ_BODY)
    add_para(doc)

    rows = lo['cal_rows']
    tbl = doc.add_table(rows=1 + len(rows), cols=4)
    _set_table_borders(tbl)
    _set_col_widths(tbl, [1.2, 9.8, 3.0, 3.5])

    # Header
    hdrs = ['№', 'Name of topics', 'Date of lesson', 'Type of activity']
    for ci, h in enumerate(hdrs):
        _cell_para(tbl.rows[0].cells[ci], h, bold=True,
                   align=WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT)

    for ri, row_num in enumerate(rows):
        row_data = data['cal_rows'].get(row_num, [str(row_num), '', '', ''])
        trow = tbl.rows[ri + 1]
        _cell_para(trow.cells[0], row_data[0], align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(trow.cells[1], row_data[1], align=WD_ALIGN_PARAGRAPH.LEFT)
        _cell_para(trow.cells[2], '', align=WD_ALIGN_PARAGRAPH.CENTER)
        _cell_para(trow.cells[3], row_data[3] if len(row_data) > 3 else '',
                   align=WD_ALIGN_PARAGRAPH.LEFT)


def build_lo_glossary(doc, lo, data):
    add_para(doc)
    add_para(doc, 'GLOSSARY',
             bold=True, size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    add_para(doc,
             'The glossary contains the main definitions, terms and concepts '
             'that reflect the core content of this learning outcome. '
             'It is arranged in alphabetical order.',
             size=SZ_BODY)
    add_para(doc)

    terms = lo['glossary_terms']
    tbl = doc.add_table(rows=1 + len(terms), cols=2)
    _set_table_borders(tbl)
    _set_col_widths(tbl, [4.5, 13.0])

    _cell_para(tbl.rows[0].cells[0], 'Term', bold=True)
    _cell_para(tbl.rows[0].cells[1], 'Definition', bold=True)

    for ri, term in enumerate(terms):
        defn = data['gloss'].get(term, '')
        trow = tbl.rows[ri + 1]
        _cell_para(trow.cells[0], term)
        _cell_para(trow.cells[1], defn)


def build_lo_lectures(doc, lo, data):
    if not lo['lecture_nums']:
        return
    add_para(doc)
    add_para(doc, 'LECTURE MATERIALS',
             bold=True, size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    add_para(doc, f'LO {lo["id"]} {lo["title"]}',
             bold=True, size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc)

    lo_id = lo['id']
    seq = 0
    pending_heading = None  # "Lecture X.Y.Z" waiting for its Topic line

    def _flush_pending():
        nonlocal pending_heading
        if pending_heading is not None:
            add_para(doc, pending_heading, bold=True, size=SZ_BODY)
            pending_heading = None

    for lnum in lo['lecture_nums']:
        for et, el in data['lectures'].get(lnum, []):
            if et == 'para':
                txt = el.text.strip()
                if txt.startswith('Lecture No.'):
                    _flush_pending()
                    seq += 1
                    pending_heading = f'Lecture {lo_id}.{seq}'
                elif pending_heading is not None and txt.startswith('Topic:'):
                    topic = txt[len('Topic:'):].strip()
                    add_para(doc, f'{pending_heading} {topic}',
                             bold=True, size=SZ_BODY)
                    pending_heading = None
                else:
                    _flush_pending()
                    if not txt:
                        add_para(doc)
                    else:
                        copy_para_to_doc(el, doc, force_size=SZ_BODY)
            elif et == 'table':
                _flush_pending()
                copy_table_to_doc(el, doc)
    _flush_pending()


def build_lo_practicals(doc, lo, data):
    label = lo['practical_label']
    add_para(doc)
    if label == 'Industrial training':
        section_heading = 'INDUSTRIAL TRAINING MATERIALS'
        old_prefix = 'Industrial Training Task No.'
        new_prefix_tmpl = 'Industrial training task {lo_id}.{seq}'
    elif label == 'Laboratory work':
        section_heading = 'LABORATORY WORK MATERIALS'
        old_prefix = 'Laboratory Work No.'
        new_prefix_tmpl = 'Laboratory work {lo_id}.{seq}'
    else:
        section_heading = 'PRACTICAL CLASS MATERIALS'
        old_prefix = 'Practical Class No.'
        new_prefix_tmpl = 'Practical work {lo_id}.{seq}'

    add_para(doc, section_heading, bold=True, size=SZ_BODY,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)

    lo_id = lo['id']
    for seq, pnum in enumerate(lo['practical_nums'], start=1):
        if label == 'Practical class':
            elems = data['practical_classes'].get(pnum, [])
        elif label == 'Laboratory work':
            elems = data['lab_works'].get(pnum, [])
        else:
            elems = data['ind_tasks'].get(pnum, [])

        new_num = new_prefix_tmpl.format(lo_id=lo_id, seq=seq)

        pending_work_heading = None

        def _flush_work():
            nonlocal pending_work_heading
            if pending_work_heading is not None:
                add_para(doc, pending_work_heading, bold=True, size=SZ_BODY)
                pending_work_heading = None

        for et, el in elems:
            if et == 'para':
                txt = el.text.strip()
                if txt.startswith(old_prefix):
                    _flush_work()
                    pending_work_heading = new_num
                elif pending_work_heading is not None and txt.startswith('Topic:'):
                    topic = txt[len('Topic:'):].strip()
                    add_para(doc, f'{pending_work_heading} {topic}',
                             bold=True, size=SZ_BODY)
                    pending_work_heading = None
                else:
                    _flush_work()
                    if not txt:
                        add_para(doc)
                    else:
                        copy_para_to_doc(el, doc, force_size=SZ_BODY)
            elif et == 'table':
                _flush_work()
                copy_table_to_doc(el, doc)
        _flush_work()


def build_lo_software(doc, lo):
    add_para(doc)
    add_para(doc,
             'LIST OF SOFTWARE AND MULTIMEDIA SUPPORT FOR TRAINING SESSIONS',
             bold=True, size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    for item in SOFTWARE_LIST:
        add_para(doc, item, size=SZ_BODY)


def build_lo_literature(doc, lo):
    add_para(doc)
    add_para(doc, 'LIST OF LITERATURE',
             bold=True, size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    add_para(doc, 'Core (fundamental) literature:', bold=True, size=SZ_BODY)
    for item in LITERATURE_CORE:
        add_para(doc, item, size=SZ_BODY)
    add_para(doc, 'Additional literature:', bold=True, size=SZ_BODY)
    for item in LITERATURE_EXTRA:
        add_para(doc, item, size=SZ_BODY)
    add_para(doc, 'Recommended learning resources:', bold=True, size=SZ_BODY)
    for item in LITERATURE_WEB:
        add_para(doc, item, size=SZ_BODY)


def build_self_study(doc, data):
    _real_page_break(doc)
    add_para(doc, 'Self-Study Topics for the Module',
             bold=True, size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    for et, el in data['self_study']:
        if et != 'para':
            continue
        txt = el.text.strip()
        # Stop before the 7.3 exam-questions sub-section
        if '7.3' in txt and 'control' in txt.lower():
            break
        # Skip original PM01 section headers that start with roman numerals
        if txt.startswith('VII.') or txt.startswith('VIII.') or txt.startswith('IX.'):
            continue
        if not txt:
            add_para(doc)
        else:
            copy_para_to_doc(el, doc, force_size=SZ_BODY)


def build_exam_questions(doc, data):
    _real_page_break(doc)
    add_para(doc, 'Exam Questions',
             bold=True, size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    in_exam = False
    for et, el in data['self_study']:
        if et != 'para':
            continue
        txt = el.text.strip()
        if '7.3' in txt and 'control' in txt.lower():
            in_exam = True
            continue   # skip the "7.3 List of…" header line itself
        if not in_exam:
            continue
        # Stop at guidelines block that follows the question list
        if 'guidelines' in txt.lower() and 'completing' in txt.lower():
            break
        if not txt:
            add_para(doc)
        else:
            copy_para_to_doc(el, doc, force_size=SZ_BODY)


# ═══════════════════════════════════════════════════════════════════════════
#  MAIN
# ═══════════════════════════════════════════════════════════════════════════

def main():
    print('Extracting PM01 content...')
    data = extract_pm01()
    print(f'  Lectures found: {sorted(data["lectures"].keys())}')
    print(f'  Practical classes: {sorted(data["practical_classes"].keys())}')
    print(f'  Lab works: {sorted(data["lab_works"].keys())}')
    print(f'  Industrial tasks: {sorted(data["ind_tasks"].keys())}')
    print(f'  Calendar rows: {sorted(data["cal_rows"].keys())}')
    print(f'  Glossary terms: {len(data["gloss"])}')

    print('Building new document...')
    doc = Document()
    set_margins(doc)

    # 1. Title page
    build_title_page(doc)

    # 2. Approval sheet
    build_approval_sheet(doc)

    # 3. Contents
    build_contents(doc)

    # 4. Introduction
    build_introduction(doc, data)

    # 5. Academic Calendar of Module (weight table)
    build_weight_table(doc)

    # 6. LO sections
    for lo in LO_LIST:
        print(f'  Building LO {lo["id"]}...')
        _real_page_break(doc)

        build_lo_calendar(doc, lo, data)
        build_lo_glossary(doc, lo, data)
        build_lo_lectures(doc, lo, data)
        build_lo_practicals(doc, lo, data)
        build_lo_software(doc, lo)
        build_lo_literature(doc, lo)

    # 7. Self-study topics
    build_self_study(doc, data)

    # 8. Exam questions
    build_exam_questions(doc, data)

    doc.save(OUTPUT_PATH)
    print(f'\nDone! Saved to: {OUTPUT_PATH}')


if __name__ == '__main__':
    main()
