# -*- coding: utf-8 -*-
"""Build a standalone, neatly formatted CONTENT (table of contents) .docx with
page numbers, following the УМКД structure rules (Times New Roman, dot leaders,
right-aligned page numbers, bold main entries, indented sub-entries).

Reads page numbers from the Word extraction (headings.txt) + the LO JSONs.
Does NOT touch УМКД_ПМ01_Готовый.docx.
"""
import io, json, os
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_TAB_LEADER
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
SCRATCH = r'C:\Users\Dias\AppData\Local\Temp\claude\C--Users-Dias-Desktop-course\767528e1-1ebd-40dc-8f49-105500f0c51f\scratchpad'
HEADINGS = os.path.join(SCRATCH, 'headings.txt')
OUT = os.path.join(HERE, '..', 'Содержание.docx')

FONT = 'Times New Roman'
SZ_H = 14
SZ = 12
ROMAN = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X']
RIGHT_TAB = int(17 * 567)  # 17 cm text width (A4 21 − L2.5 − R1.5) in twips

# ── reconstruct (toc_label, real_heading, is_main) ───────────────────────────
los = [json.load(open(os.path.join(HERE, f'lo_{k}.json'), encoding='utf-8'))
       for k in ['1_1', '1_2', '1_3', '1_4', '1_5', '1_6', '1_7']]
SECTION_MAT = {'practical': 'PRACTICAL CLASS MATERIALS',
               'lab': 'LABORATORY WORK MATERIALS',
               'industrial': 'INDUSTRIAL TRAINING MATERIALS'}

pairs = []  # (toc_label, real_heading, is_main)
pairs.append(('INTRODUCTION', 'INTRODUCTION', True))
pairs.append(('I. ACADEMIC CALENDAR OF MODULE', 'I. ACADEMIC CALENDAR OF MODULE', True))
for idx, lo in enumerate(los):
    rn, lid = ROMAN[idx + 1], lo['id']
    has_lec = any(l['kind'] == 'lecture' for l in lo['lessons'])
    others = [l for l in lo['lessons'] if l['kind'] != 'lecture']
    kind = others[0]['kind'] if others else 'practical'
    pairs.append((f'{rn}. LO {lid} {lo["title"]}', f'{rn}. LO {lid} {lo["title"]}', True))
    pairs.append((f'1. LO {lid} Academic calendar', '1. ACADEMIC CALENDAR', False))
    pairs.append((f'2. LO {lid} Glossary', '2. GLOSSARY', False))
    if has_lec:
        pairs.append((f'3. LO {lid} Lecture materials', '3. LECTURE MATERIALS', False))
    pairs.append((f'4. LO {lid} Practical / training materials', f'4. {SECTION_MAT[kind]}', False))
    pairs.append((f'5. LO {lid} List of software and multimedia support',
                  '5. LIST OF SOFTWARE AND MULTIMEDIA SUPPORT FOR TRAINING SESSIONS', False))
    pairs.append((f'6. LO {lid} List of literature', '6. LIST OF LITERATURE', False))
pairs.append(('IX. Self-Study Topics for the Module', 'IX. SELF-STUDY TOPICS FOR THE MODULE', True))
pairs.append(('X. Exam Questions', 'X. EXAM QUESTIONS', True))

# ── page numbers from the Word extraction ────────────────────────────────────
rows = []
for line in io.open(HEADINGS, encoding='utf-8'):
    if not line.strip():
        continue
    pg, txt = line.rstrip('\n').split('\t', 1)
    rows.append((int(pg), txt.strip()))
intro_idx = [i for i, (_, t) in enumerate(rows) if t == 'INTRODUCTION']
real = rows[intro_idx[1]:]
assert len(real) == len(pairs), f'count mismatch {len(real)} vs {len(pairs)}'
for (toc, expected, _), (pg, got) in zip(pairs, real):
    assert got == expected, f'misalign: {got!r} != {expected!r}'
entries = [(toc, pg, is_main) for (toc, _, is_main), (pg, _) in zip(pairs, real)]

# ── build the document ───────────────────────────────────────────────────────
doc = Document()
sec = doc.sections[0]
sectPr = sec._sectPr
pgMar = sectPr.find(qn('w:pgMar'))
if pgMar is None:
    pgMar = OxmlElement('w:pgMar')
    sectPr.append(pgMar)
pgMar.set(qn('w:top'), str(int(20 * 56.7)))
pgMar.set(qn('w:bottom'), str(int(20 * 56.7)))
pgMar.set(qn('w:left'), str(int(25 * 56.7)))
pgMar.set(qn('w:right'), str(int(15 * 56.7)))


def _run(r, bold=False, size=SZ):
    r.font.name = FONT
    r.font.size = Pt(size)
    r.font.bold = bold
    rpr = r._r.get_or_add_rPr()
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = OxmlElement('w:rFonts')
        rpr.append(rf)
    rf.set(qn('w:ascii'), FONT)
    rf.set(qn('w:hAnsi'), FONT)
    rf.set(qn('w:cs'), FONT)


# heading
h = doc.add_paragraph()
h.alignment = WD_ALIGN_PARAGRAPH.CENTER
h.paragraph_format.space_after = Pt(12)
h.paragraph_format.line_spacing = 1.0
_run(h.add_run('CONTENT'), bold=True, size=SZ_H)

for toc, pg, is_main in entries:
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = WD_ALIGN_PARAGRAPH.LEFT
    pf.line_spacing = 1.0
    pf.space_after = Pt(6 if is_main else 2)
    if not is_main:
        pf.left_indent = Pt(24)
    # right tab with dotted leader at the text-area right edge
    pf.tab_stops.add_tab_stop(Pt(RIGHT_TAB / 20.0), WD_TAB_ALIGNMENT.RIGHT,
                              WD_TAB_LEADER.DOTS)
    _run(p.add_run(toc), bold=is_main)
    _run(p.add_run('\t' + str(pg)), bold=is_main)

doc.save(OUT)
print('saved:', os.path.abspath(OUT))
print('entries:', len(entries))
