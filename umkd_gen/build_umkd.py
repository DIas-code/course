"""
Build the complete УМКД ПМ 01 docx from structured content.

Reads:
  - module.json        : module-level metadata (title page, intro, weights, software, literature, exam, self-study)
  - lo_1_1.json ... lo_1_7.json : per-LO definitions + lessons (full English content)

Produces:
  - ../УМКД_ПМ01_Готовый.docx

Formatting follows the Бекет А. ПМ10 example (эталон) and the ТИПО-2025
methodical recommendations: Times New Roman, body 12 pt, justified, single
spacing, A4 with margins left 25 / top 20 / right 15 / bottom 20 mm, centered
page numbers (10 pt) in the footer.
"""

import os
import sys
import json

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
DEFAULT_OUT = 'УМКД_ПМ01_Готовый.docx'
DEFAULT_LO_FILES = ['1_1', '1_2', '1_3', '1_4', '1_5', '1_6', '1_7']

FONT = 'Times New Roman'
SZ_TITLE = 14
SZ_BODY = 12
SZ_SMALL = 10

ROMAN = ['I', 'II', 'III', 'IV', 'V', 'VI', 'VII', 'VIII', 'IX', 'X',
         'XI', 'XII', 'XIII', 'XIV', 'XV']

# ── low-level helpers ─────────────────────────────────────────────────────────

def _set_run(run, bold=False, size=SZ_BODY, italic=False, font=FONT):
    run.font.name = font
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    # ensure east-asian / cs fonts also TNR so the size sticks everywhere
    rpr = run._r.get_or_add_rPr()
    rfonts = rpr.find(qn('w:rFonts'))
    if rfonts is None:
        rfonts = OxmlElement('w:rFonts')
        rpr.append(rfonts)
    rfonts.set(qn('w:ascii'), font)
    rfonts.set(qn('w:hAnsi'), font)
    rfonts.set(qn('w:cs'), font)


def add_para(doc, text='', bold=False, size=SZ_BODY, italic=False,
             align=WD_ALIGN_PARAGRAPH.JUSTIFY, space_before=0, space_after=0,
             left_indent=None, first_line=None):
    p = doc.add_paragraph()
    pf = p.paragraph_format
    pf.alignment = align
    pf.space_before = Pt(space_before)
    pf.space_after = Pt(space_after)
    pf.line_spacing = 1.0
    if left_indent is not None:
        pf.left_indent = Pt(left_indent)
    if first_line is not None:
        pf.first_line_indent = Pt(first_line)
    if text:
        r = p.add_run(text)
        _set_run(r, bold=bold, size=size, italic=italic)
    return p


def heading(doc, text, size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.CENTER,
            space_before=6, space_after=6):
    return add_para(doc, text, bold=True, size=size, align=align,
                    space_before=space_before, space_after=space_after)


def page_break(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run()
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    run._r.append(br)


def set_margins(doc):
    # left 25, top 20, right 15, bottom 20 mm  (1 mm ≈ 56.7 twips)
    for section in doc.sections:
        sectPr = section._sectPr
        pgMar = sectPr.find(qn('w:pgMar'))
        if pgMar is None:
            pgMar = OxmlElement('w:pgMar')
            sectPr.append(pgMar)
        pgMar.set(qn('w:top'), str(int(20 * 56.7)))
        pgMar.set(qn('w:bottom'), str(int(20 * 56.7)))
        pgMar.set(qn('w:left'), str(int(25 * 56.7)))
        pgMar.set(qn('w:right'), str(int(15 * 56.7)))


def add_page_numbers(doc):
    """Centered page number (10 pt) in the footer of every section."""
    for section in doc.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        p.text = ''
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        _set_run(run, size=SZ_SMALL)
        fldStart = OxmlElement('w:fldChar')
        fldStart.set(qn('w:fldCharType'), 'begin')
        instr = OxmlElement('w:instrText')
        instr.set(qn('xml:space'), 'preserve')
        instr.text = 'PAGE'
        fldEnd = OxmlElement('w:fldChar')
        fldEnd.set(qn('w:fldCharType'), 'end')
        run._r.append(fldStart)
        run._r.append(instr)
        run._r.append(fldEnd)


# ── table helpers ─────────────────────────────────────────────────────────────

def _set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr
    tblBorders = OxmlElement('w:tblBorders')
    for side in ('top', 'left', 'bottom', 'right', 'insideH', 'insideV'):
        el = OxmlElement(f'w:{side}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), '4')
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), '000000')
        tblBorders.append(el)
    existing = tblPr.find(qn('w:tblBorders'))
    if existing is not None:
        tblPr.remove(existing)
    tblPr.append(tblBorders)


def _set_col_widths(table, widths_cm):
    for row in table.rows:
        for ci, cell in enumerate(row.cells):
            if ci < len(widths_cm):
                tc = cell._tc
                tcPr = tc.get_or_add_tcPr()
                tcW = tcPr.find(qn('w:tcW'))
                if tcW is None:
                    tcW = OxmlElement('w:tcW')
                    tcPr.append(tcW)
                tcW.set(qn('w:w'), str(int(widths_cm[ci] * 567)))
                tcW.set(qn('w:type'), 'dxa')


def cell_text(cell, text, bold=False, size=SZ_BODY,
              align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ''
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1.0
    r = p.add_run(text if text is not None else '')
    _set_run(r, bold=bold, size=size)


# ── section builders ──────────────────────────────────────────────────────────

def build_title_page(doc, m):
    add_para(doc, m['college'], bold=True, size=SZ_TITLE,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)

    tbl = doc.add_table(rows=1, cols=2)
    _set_col_widths(tbl, [9.0, 8.0])
    left, right = tbl.rows[0].cells
    left.text = ''
    right.text = ''
    rp = right.paragraphs[0]
    for i, line in enumerate(m['approved_block']):
        p = rp if i == 0 else right.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(line)
        _set_run(r, bold=(i == 0), size=SZ_BODY)

    for _ in range(3):
        add_para(doc)
    add_para(doc, 'Educational-methodological complex of the module',
             bold=True, size=SZ_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    add_para(doc, m['module_title'], bold=True, size=SZ_TITLE,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    add_para(doc, m['specialty'], size=SZ_TITLE, align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc, m['qualification'], size=SZ_TITLE,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    for _ in range(5):
        add_para(doc)
    add_para(doc, m['place_year'], bold=True, size=SZ_TITLE,
             align=WD_ALIGN_PARAGRAPH.CENTER)


def build_approval_sheet(doc, m):
    page_break(doc)
    add_para(doc, m['college'], bold=True, size=SZ_TITLE,
             align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    heading(doc, 'APPROVAL SHEET', size=SZ_TITLE)
    add_para(doc)
    for line in m['approval_sheet']:
        if line == '':
            add_para(doc)
        else:
            add_para(doc, line, size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.LEFT)


def build_contents(doc, los):
    page_break(doc)
    heading(doc, 'CONTENT', size=SZ_BODY)
    add_para(doc)

    def toc(text, bold=False, indent=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        if indent:
            p.paragraph_format.left_indent = Pt(22)
        r = p.add_run(text)
        _set_run(r, bold=bold, size=SZ_BODY)

    toc('INTRODUCTION', bold=True)
    toc('I. ACADEMIC CALENDAR OF MODULE', bold=True)
    for idx, lo in enumerate(los):
        rn = ROMAN[idx + 1]
        toc(f'{rn}. LO {lo["id"]} {lo["title"]}', bold=True)
        toc(f'1. LO {lo["id"]} Academic calendar', indent=True)
        toc(f'2. LO {lo["id"]} Glossary', indent=True)
        if any(l['kind'] == 'lecture' for l in lo['lessons']):
            toc(f'3. LO {lo["id"]} Lecture materials', indent=True)
        if any(l['kind'] != 'lecture' for l in lo['lessons']):
            toc(f'4. LO {lo["id"]} Practical / training materials', indent=True)
        toc(f'5. LO {lo["id"]} List of software and multimedia support', indent=True)
        toc(f'6. LO {lo["id"]} List of literature', indent=True)
    n = len(los) + 2
    toc(f'{ROMAN[n-1]}. Self-Study Topics for the Module', bold=True)
    toc(f'{ROMAN[n]}. Exam Questions', bold=True)


def build_introduction(doc, m):
    page_break(doc)
    heading(doc, 'INTRODUCTION', size=SZ_BODY)
    add_para(doc)
    for para in m['introduction']:
        add_para(doc, para, size=SZ_BODY, first_line=18)


def build_weight_table(doc, m, los):
    page_break(doc)
    heading(doc, 'I. ACADEMIC CALENDAR OF MODULE', size=SZ_BODY)
    add_para(doc)
    tbl = doc.add_table(rows=1 + len(los) + 1, cols=3)
    _set_table_borders(tbl)
    _set_col_widths(tbl, [2.5, 12.0, 2.5])
    cell_text(tbl.rows[0].cells[0], m['module_code'], bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(tbl.rows[0].cells[1], m['module_name'], bold=True)
    cell_text(tbl.rows[0].cells[2], 'Weight, %', bold=True,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    for i, lo in enumerate(los):
        row = tbl.rows[i + 1]
        cell_text(row.cells[0], f'LO {lo["id"]}', align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row.cells[1], lo['title'])
        cell_text(row.cells[2], lo['weight'], align=WD_ALIGN_PARAGRAPH.CENTER)
    last = tbl.rows[-1]
    cell_text(last.cells[0], '', align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_text(last.cells[1], 'Intermediate certification (exam)')
    cell_text(last.cells[2], m['exam_weight'], align=WD_ALIGN_PARAGRAPH.CENTER)


def build_lo_calendar(doc, m, lo, roman):
    heading(doc, f'{roman}. LO {lo["id"]} {lo["title"]}', size=SZ_BODY)
    add_para(doc)
    heading(doc, '1. ACADEMIC CALENDAR', size=SZ_BODY,
            align=WD_ALIGN_PARAGRAPH.CENTER)
    add_para(doc)
    add_para(doc, f'LO {lo["id"]} {lo["title"]}', bold=True, size=SZ_BODY,
             align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, f'Instructor: {m["instructor"]}', align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, m['specialty'], align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, f'Total hours: {lo["hours"]} ({len(lo["lessons"])} lessons '
                  f'of 2 academic hours each)', align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc)

    tbl = doc.add_table(rows=1 + len(lo['lessons']), cols=4)
    _set_table_borders(tbl)
    _set_col_widths(tbl, [1.2, 10.3, 2.5, 3.5])
    hdrs = ['№', 'Name of topics', 'Date of lesson', 'Type of activity']
    for ci, h in enumerate(hdrs):
        cell_text(tbl.rows[0].cells[ci], h, bold=True,
                  align=WD_ALIGN_PARAGRAPH.CENTER if ci != 1 else WD_ALIGN_PARAGRAPH.LEFT)
    for i, les in enumerate(lo['lessons']):
        row = tbl.rows[i + 1]
        cell_text(row.cells[0], str(i + 1), align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row.cells[1], les['title'])
        cell_text(row.cells[2], '', align=WD_ALIGN_PARAGRAPH.CENTER)
        cell_text(row.cells[3], les['type_en'])


def build_lo_glossary(doc, lo):
    add_para(doc)
    heading(doc, '2. GLOSSARY', size=SZ_BODY)
    add_para(doc)
    add_para(doc, 'The glossary contains the main definitions, terms and '
                  'concepts that reflect the core content of this learning '
                  'outcome. The terms are arranged in alphabetical order.',
             size=SZ_BODY, first_line=18)
    add_para(doc)
    terms = lo['glossary']
    tbl = doc.add_table(rows=1 + len(terms), cols=2)
    _set_table_borders(tbl)
    _set_col_widths(tbl, [4.5, 13.0])
    cell_text(tbl.rows[0].cells[0], 'Term', bold=True)
    cell_text(tbl.rows[0].cells[1], 'Definition', bold=True)
    for i, (term, defn) in enumerate(terms):
        row = tbl.rows[i + 1]
        cell_text(row.cells[0], term, bold=True)
        cell_text(row.cells[1], defn)


def _render_lecture(doc, les, label):
    add_para(doc, f'{label} {les["title"]}', bold=True, size=SZ_BODY,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6)
    add_para(doc)
    add_para(doc, 'Content', bold=True, size=SZ_BODY,
             align=WD_ALIGN_PARAGRAPH.LEFT)
    for item in les.get('outline', []):
        add_para(doc, item, size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.LEFT,
                 left_indent=18)
    add_para(doc)
    for para in les.get('body', []):
        if para.startswith('• '):
            add_para(doc, para, size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.LEFT,
                     left_indent=18)
        elif para.startswith('## '):
            add_para(doc, para[3:], bold=True, size=SZ_BODY,
                     align=WD_ALIGN_PARAGRAPH.LEFT)
        else:
            add_para(doc, para, size=SZ_BODY, first_line=18)
    add_para(doc)
    add_para(doc, 'Control Questions', bold=True, size=SZ_BODY,
             align=WD_ALIGN_PARAGRAPH.LEFT)
    for i, q in enumerate(les.get('questions', []), 1):
        add_para(doc, f'{i}. {q}', size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc)


def _render_practical(doc, les, label):
    add_para(doc, f'{label} {les["title"]}', bold=True, size=SZ_BODY,
             align=WD_ALIGN_PARAGRAPH.LEFT, space_before=6)
    add_para(doc)
    if les.get('objective'):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r1 = p.add_run('Objective: ')
        _set_run(r1, bold=True, size=SZ_BODY)
        r2 = p.add_run(les['objective'])
        _set_run(r2, size=SZ_BODY)
    if les.get('theory'):
        add_para(doc, 'Theoretical background', bold=True, size=SZ_BODY,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
        for para in les['theory']:
            if para.startswith('• '):
                add_para(doc, para, size=SZ_BODY, left_indent=18,
                         align=WD_ALIGN_PARAGRAPH.LEFT)
            else:
                add_para(doc, para, size=SZ_BODY, first_line=18)
    if les.get('materials'):
        add_para(doc, 'Equipment and materials', bold=True, size=SZ_BODY,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
        for it in les['materials']:
            add_para(doc, f'– {it}', size=SZ_BODY, left_indent=18,
                     align=WD_ALIGN_PARAGRAPH.LEFT)
    if les.get('procedure'):
        add_para(doc, 'Procedure', bold=True, size=SZ_BODY,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
        for i, step in enumerate(les['procedure'], 1):
            add_para(doc, f'{i}. {step}', size=SZ_BODY, first_line=0,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    if les.get('tasks'):
        add_para(doc, 'Tasks for independent completion', bold=True,
                 size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.LEFT)
        for i, t in enumerate(les['tasks'], 1):
            add_para(doc, f'{i}. {t}', size=SZ_BODY,
                     align=WD_ALIGN_PARAGRAPH.JUSTIFY)
    if les.get('conclusion'):
        p = doc.add_paragraph()
        p.paragraph_format.line_spacing = 1.0
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        r1 = p.add_run('Conclusion: ')
        _set_run(r1, bold=True, size=SZ_BODY)
        r2 = p.add_run(les['conclusion'])
        _set_run(r2, size=SZ_BODY)
    if les.get('questions'):
        add_para(doc, 'Control Questions', bold=True, size=SZ_BODY,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
        for i, q in enumerate(les['questions'], 1):
            add_para(doc, f'{i}. {q}', size=SZ_BODY,
                     align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc)


PRACTICAL_LABELS = {
    'practical': 'Practical work',
    'lab': 'Laboratory work',
    'industrial': 'Industrial training task',
}
PRACTICAL_SECTIONS = {
    'practical': 'PRACTICAL CLASS MATERIALS',
    'lab': 'LABORATORY WORK MATERIALS',
    'industrial': 'INDUSTRIAL TRAINING MATERIALS',
}


def build_lo_lessons(doc, lo):
    lectures = [l for l in lo['lessons'] if l['kind'] == 'lecture']
    others = [l for l in lo['lessons'] if l['kind'] != 'lecture']

    if lectures:
        add_para(doc)
        heading(doc, '3. LECTURE MATERIALS', size=SZ_BODY)
        add_para(doc)
        add_para(doc, f'LO {lo["id"]} {lo["title"]}', bold=True, size=SZ_BODY,
                 align=WD_ALIGN_PARAGRAPH.LEFT)
        add_para(doc)
        for les in lectures:
            label = f'Lecture {lo["id"]}.{les["pos"]}'
            _render_lecture(doc, les, label)

    if others:
        # group consecutive by kind, but render under one heading per kind present
        kinds_present = []
        for l in others:
            if l['kind'] not in kinds_present:
                kinds_present.append(l['kind'])
        add_para(doc)
        # primary heading uses the first kind, but we render each kind section
        for ki, kind in enumerate(kinds_present):
            heading(doc, f'4. {PRACTICAL_SECTIONS[kind]}', size=SZ_BODY)
            add_para(doc)
            add_para(doc, f'LO {lo["id"]} {lo["title"]}', bold=True,
                     size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.LEFT)
            add_para(doc)
            for les in [l for l in others if l['kind'] == kind]:
                label = f'{PRACTICAL_LABELS[kind]} {lo["id"]}.{les["pos"]}'
                _render_practical(doc, les, label)


def build_lo_software(doc, m):
    add_para(doc)
    heading(doc, '5. LIST OF SOFTWARE AND MULTIMEDIA SUPPORT FOR TRAINING SESSIONS',
            size=SZ_BODY)
    add_para(doc)
    for it in m['software']:
        add_para(doc, it, size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.LEFT)


def build_lo_literature(doc, m):
    add_para(doc)
    heading(doc, '6. LIST OF LITERATURE', size=SZ_BODY)
    add_para(doc)
    add_para(doc, 'Core (fundamental) literature:', bold=True, size=SZ_BODY,
             align=WD_ALIGN_PARAGRAPH.LEFT)
    for it in m['literature']['core']:
        add_para(doc, it, size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, 'Additional literature:', bold=True, size=SZ_BODY,
             align=WD_ALIGN_PARAGRAPH.LEFT)
    for it in m['literature']['extra']:
        add_para(doc, it, size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.LEFT)
    add_para(doc, 'Internet resources:', bold=True, size=SZ_BODY,
             align=WD_ALIGN_PARAGRAPH.LEFT)
    for it in m['literature']['web']:
        add_para(doc, it, size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.LEFT)


def build_self_study(doc, m, n):
    page_break(doc)
    heading(doc, f'{ROMAN[n-1]}. SELF-STUDY TOPICS FOR THE MODULE', size=SZ_BODY)
    add_para(doc)
    add_para(doc, m['self_study_intro'], size=SZ_BODY, first_line=18)
    add_para(doc)
    for i, t in enumerate(m['self_study'], 1):
        add_para(doc, f'{i}. {t}', size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.LEFT)


def build_exam(doc, m, n):
    page_break(doc)
    heading(doc, f'{ROMAN[n]}. EXAM QUESTIONS', size=SZ_BODY)
    add_para(doc)
    add_para(doc, m['exam_intro'], size=SZ_BODY, first_line=18)
    add_para(doc)
    for i, q in enumerate(m['exam_questions'], 1):
        add_para(doc, f'{i}. {q}', size=SZ_BODY, align=WD_ALIGN_PARAGRAPH.LEFT)


# ── main ──────────────────────────────────────────────────────────────────────

def load_json(name):
    with open(os.path.join(HERE, name), encoding='utf-8') as f:
        return json.load(f)


def main():
    module_file = sys.argv[1] if len(sys.argv) > 1 else 'module.json'
    m = load_json(module_file)
    lo_files = m.get('lo_files', DEFAULT_LO_FILES)
    out_name = m.get('out_name', DEFAULT_OUT)
    out_path = os.path.join(HERE, '..', out_name)
    los = []
    for lo_id in lo_files:
        lo = load_json(f'lo_{lo_id}.json')
        # assign positions within LO
        for i, les in enumerate(lo['lessons'], 1):
            les['pos'] = i
        los.append(lo)

    doc = Document()
    set_margins(doc)

    build_title_page(doc, m)
    build_approval_sheet(doc, m)
    build_contents(doc, los)
    build_introduction(doc, m)
    build_weight_table(doc, m, los)

    for idx, lo in enumerate(los):
        page_break(doc)
        roman = ROMAN[idx + 1]
        build_lo_calendar(doc, m, lo, roman)
        build_lo_glossary(doc, lo)
        build_lo_lessons(doc, lo)
        build_lo_software(doc, m)
        build_lo_literature(doc, m)

    n = len(los) + 2
    build_self_study(doc, m, n)
    build_exam(doc, m, n)

    add_page_numbers(doc)
    doc.save(out_path)
    total = sum(len(lo['lessons']) for lo in los)
    print(f'Saved: {os.path.abspath(out_path)}')
    print(f'LOs: {len(los)}, total lessons: {total}')
    for lo in los:
        lec = sum(1 for l in lo["lessons"] if l["kind"] == "lecture")
        oth = len(lo["lessons"]) - lec
        print(f'  LO {lo["id"]}: {len(lo["lessons"])} lessons '
              f'({lec} lectures, {oth} practical/training), {lo["hours"]}')


if __name__ == '__main__':
    main()
